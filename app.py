import os
import re
import io
import uuid
import sqlite3
import smtplib
import requests
import qrcode
from email.message import EmailMessage
from flask import Flask, request, jsonify, send_file, session, redirect
from dotenv import load_dotenv
from openai import OpenAI
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address
from paypalcheckoutsdk.core import PayPalHttpClient, SandboxEnvironment, LiveEnvironment
from paypalcheckoutsdk.orders import OrdersCreateRequest
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors

# ------------------ LOAD ENV ------------------
load_dotenv()

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
EMAIL_ADDRESS = os.getenv("EMAIL_ADDRESS")
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD")
PAYPAL_CLIENT_ID = os.getenv("PAYPAL_CLIENT_ID")
PAYPAL_CLIENT_SECRET = os.getenv("PAYPAL_CLIENT_SECRET")
PAYPAL_MODE = os.getenv("PAYPAL_MODE", "sandbox")
ADMIN_PASSWORD = os.getenv("ADMIN_PASSWORD")
RECAPTCHA_SECRET_KEY = os.getenv("RECAPTCHA_SECRET_KEY")

# ------------------ APP INIT ------------------
app = Flask(__name__)
app.secret_key = ADMIN_PASSWORD
limiter = Limiter(get_remote_address, app=app)

client = OpenAI(api_key=OPENAI_API_KEY)

if PAYPAL_MODE == "live":
    environment = LiveEnvironment(PAYPAL_CLIENT_ID, PAYPAL_CLIENT_SECRET)
else:
    environment = SandboxEnvironment(PAYPAL_CLIENT_ID, PAYPAL_CLIENT_SECRET)

paypal_client = PayPalHttpClient(environment)

if not os.path.exists("generated_files"):
    os.makedirs("generated_files")

# ------------------ DATABASE ------------------
def init_db():
    conn = sqlite3.connect("orders.db")
    c = conn.cursor()
    c.execute("""
        CREATE TABLE IF NOT EXISTS orders (
            id TEXT PRIMARY KEY,
            email TEXT,
            resume TEXT,
            job TEXT,
            template TEXT,
            status TEXT,
            pdf_path TEXT,
            docx_path TEXT,
            created_at DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    """)
    conn.commit()
    conn.close()

init_db()

# ------------------ UTILITIES ------------------

def extract_candidate_name(text):
    lines = text.strip().split("\n")
    return lines[0] if lines else "Candidate"

def verify_captcha(token):
    r = requests.post(
        "https://www.google.com/recaptcha/api/siteverify",
        data={"secret": RECAPTCHA_SECRET_KEY, "response": token}
    )
    return r.json().get("success", False)

def calculate_match(resume, job):
    resume_words = set(re.findall(r'\w+', resume.lower()))
    job_words = set(re.findall(r'\w+', job.lower()))
    if not job_words:
        return 0, 0
    match = len(resume_words & job_words)
    percent = round((match / len(job_words)) * 100)
    ats = min(100, round(percent * 0.7 + 30))
    return percent, ats

# ------------------ AI ------------------

def generate_resume(resume, job):
    prompt = f"""
Rewrite and optimize this resume for the job description.
Include measurable achievements.
Generate a professional summary.
Generate a tailored cover letter.

RESUME:
{resume}

JOB:
{job}
"""
    response = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role":"user","content":prompt}]
    )
    return response.choices[0].message.content

# ------------------ PDF ------------------

def generate_pdf(content, order_id, template):
    filename = f"generated_files/{order_id}.pdf"
    doc = SimpleDocTemplate(filename, pagesize=letter)
    elements = []
    styles = getSampleStyleSheet()

    elements.append(Paragraph("Resume Optimization Report", styles["Title"]))
    elements.append(Spacer(1, 12))

    for line in content.split("\n"):
        elements.append(Paragraph(line, styles["Normal"]))
        elements.append(Spacer(1, 6))

    doc.build(elements)
    return filename

# ------------------ DOCX ------------------

def generate_docx(content, order_id):
    filename = f"generated_files/{order_id}.docx"
    document = Document()

    document.add_heading("Resume Optimization Report", 0)

    for line in content.split("\n"):
        document.add_paragraph(line)

    document.save(filename)
    return filename

# ------------------ EMAIL ------------------

def send_email(to_email, pdf_path, docx_path):
    msg = EmailMessage()
    msg["Subject"] = "Your Optimized Resume"
    msg["From"] = EMAIL_ADDRESS
    msg["To"] = to_email
    msg.set_content("Your optimized resume is attached.")

    with open(pdf_path, "rb") as f:
        msg.add_attachment(f.read(), maintype="application", subtype="pdf", filename="resume.pdf")

    with open(docx_path, "rb") as f:
        msg.add_attachment(f.read(), maintype="application",
                           subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
                           filename="resume.docx")

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
        smtp.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
        smtp.send_message(msg)

# ------------------ ROUTES ------------------

@app.route("/")
def home():
    return """
    <h2>AI Resume Optimizer</h2>
    <form method="POST" action="/create-order" enctype="multipart/form-data">
        <input type="email" name="email" placeholder="Email" required><br>
        <input type="file" name="resume" required><br>
        <textarea name="job" placeholder="Job Description" required></textarea><br>
        <select name="template">
            <option value="standard">Standard</option>
            <option value="executive">Executive</option>
            <option value="minimal">Minimal</option>
        </select><br>
        <button type="submit">Pay with PayPal</button>
    </form>
    """

@app.route("/create-order", methods=["POST"])
@limiter.limit("5 per minute")
def create_order():
    email = request.form["email"]
    job = request.form["job"]
    template = request.form["template"]
    resume_file = request.files["resume"]
    resume_text = resume_file.read().decode("utf-8")

    request_order = OrdersCreateRequest()
    request_order.request_body({
        "intent":"CAPTURE",
        "purchase_units":[{"amount":{"currency_code":"USD","value":"19.00"}}]
    })

    response = paypal_client.execute(request_order)
    order_id = response.result.id

    conn = sqlite3.connect("orders.db")
    c = conn.cursor()
    c.execute("INSERT INTO orders VALUES (?, ?, ?, ?, ?, ?, ?, ?, datetime('now'))",
              (order_id, email, resume_text, job, template, "pending", "", ""))
    conn.commit()
    conn.close()

    approve_link = next(link.href for link in response.result.links if link.rel=="approve")
    return redirect(approve_link)

@app.route("/paypal-webhook", methods=["POST"])
def paypal_webhook():
    event = request.json
    if event.get("event_type") == "CHECKOUT.ORDER.APPROVED":
        order_id